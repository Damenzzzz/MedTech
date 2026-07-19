from __future__ import annotations

import hashlib
import json
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from app.rag.schemas import ProtocolDocument
from app.rag.text import extract_icd_codes, normalize_text


SUPPORTED_SUFFIXES = {".json", ".jsonl", ".txt", ".md", ".pdf", ".docx", ".zip"}


def load_corpus(corpus_path: Path) -> list[ProtocolDocument]:
    files = collect_files(corpus_path)
    documents: list[ProtocolDocument] = []
    for file_path in files:
        documents.extend(load_file(file_path))
    return [doc for doc in documents if doc.text.strip()]


def collect_files(path: Path) -> list[Path]:
    if not path.exists():
        return []
    if path.is_file():
        return [path] if path.suffix.lower() in SUPPORTED_SUFFIXES else []
    return sorted(p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES)


def load_file(path: Path) -> list[ProtocolDocument]:
    suffix = path.suffix.lower()
    if suffix == ".zip":
        return load_zip(path)
    if suffix == ".json":
        return load_json(path)
    if suffix == ".jsonl":
        return load_jsonl(path)
    text = extract_text(path)
    if not text:
        return []
    return [document_from_text(path, text)]


def load_zip(path: Path) -> list[ProtocolDocument]:
    documents: list[ProtocolDocument] = []
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(path) as zf:
            zf.extractall(tmp_path)
        for nested in collect_files(tmp_path):
            documents.extend(load_file(nested))
    return documents


def load_json(path: Path) -> list[ProtocolDocument]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return [document_from_mapping(path, item, i) for i, item in enumerate(data) if isinstance(item, dict)]
    if isinstance(data, dict):
        if "text" in data or "content" in data:
            return [document_from_mapping(path, data, 0)]
        documents: list[ProtocolDocument] = []
        for i, value in enumerate(data.values()):
            if isinstance(value, dict):
                documents.append(document_from_mapping(path, value, i))
            elif isinstance(value, str):
                documents.append(document_from_text(path, value, suffix=str(i)))
        return documents
    return []


def load_jsonl(path: Path) -> list[ProtocolDocument]:
    documents: list[ProtocolDocument] = []
    for i, line in enumerate(path.read_text(encoding="utf-8").splitlines()):
        if not line.strip():
            continue
        item = json.loads(line)
        if isinstance(item, dict):
            documents.append(document_from_mapping(path, item, i))
    return documents


def document_from_mapping(path: Path, item: dict[str, Any], index: int) -> ProtocolDocument:
    text = normalize_text(str(item.get("text") or item.get("content") or item.get("body") or ""))
    title = str(item.get("title") or item.get("name") or item.get("source_file") or path.stem)
    source_file = str(item.get("source_file") or item.get("file") or path.name)
    protocol_id = str(item.get("protocol_id") or item.get("id") or stable_id(source_file, str(index), title))
    raw_codes = item.get("icd_codes") or item.get("icd10") or item.get("icd10_codes") or []
    if isinstance(raw_codes, str):
        raw_codes = extract_icd_codes(raw_codes)
    icd_codes = list(dict.fromkeys([str(code).upper() for code in raw_codes] + extract_icd_codes(text)))
    metadata = {k: v for k, v in item.items() if k not in {"text", "content", "body"}}
    return ProtocolDocument(
        protocol_id=protocol_id,
        source_file=source_file,
        title=title,
        icd_codes=icd_codes,
        text=text,
        metadata=metadata,
    )


def document_from_text(path: Path, text: str, suffix: str = "0") -> ProtocolDocument:
    cleaned = normalize_text(text)
    return ProtocolDocument(
        protocol_id=stable_id(path.name, suffix, cleaned[:120]),
        source_file=path.name,
        title=path.stem,
        icd_codes=extract_icd_codes(cleaned),
        text=cleaned,
        metadata={"path": str(path)},
    )


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    return ""


def stable_id(*parts: str) -> str:
    digest = hashlib.sha1("|".join(parts).encode("utf-8")).hexdigest()[:12]
    return f"p_{digest}"
